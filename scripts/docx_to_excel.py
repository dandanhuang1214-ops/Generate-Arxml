"""DOCX standard document → Excel template converter.

Two modes:
  1. create-docx:  Generate a standard DOCX template from the built-in example data.
  2. convert:      Parse a structured DOCX and produce an Excel workbook.

Usage:
  python scripts/docx_to_excel.py create-docx --output docs/horn_standard.docx
  python scripts/docx_to_excel.py convert --input docs/horn_standard.docx --output data/input/my_template.xlsx
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

# Ensure src is on path for imports
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font as XlFont, PatternFill
from openpyxl.worksheet.datavalidation import DataValidation

from arxml_codegen.excel.template import EXAMPLE_ROWS, SHEETS


# ── DOCX generation ──────────────────────────────────────────────

def create_docx_template(output: Path) -> None:
    """Generate a comprehensive standard document with the horn example."""
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A3 landscape for wide tables
    section.page_height = Cm(21.0)

    # ── Styles ──
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Title ──
    title = doc.add_heading("AUTOSAR SWC 设计标准文档", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "本文档是 ARXML 代码生成工具的标准输入格式。"
        "各节表格与 Excel 模板 Sheet 一一对应，填写完成后可通过 docx_to_excel.py 转换为 Excel 工作簿。"
    )
    doc.add_paragraph("项目：Horn Control（喇叭控制） | AUTOSAR 版本：4.3.0 | 根包：/HORN_CTRL")

    # ── Table of Contents placeholder ──
    doc.add_heading("目录", level=1)
    for name in SHEETS:
        doc.add_paragraph(f"  • {name}", style="List Bullet")

    # ── Section groups ──
    _add_section(doc, "项目配置", "ProjectConfig",
                 "项目级配置参数，Key-Value 形式。AutosarVersion 决定生成的 schema 版本，"
                 "RootPackage 是项目的顶层包路径。")
    _add_section(doc, "组件定义", "Components",
                 "定义所有 SWC 类型。ComponentKind 为 Application（原子 SWC）或 Composition（组合 SWC）。"
                 "InternalBehaviorName 和 ImplementationName 用于 DaVinci Developer 兼容。")
    _add_section(doc, "组件原型（实例化）", "ComponentPrototypes",
                 "在 Composition 内实例化组件类型。PrototypeName 是实例名，ComponentTypeName 引用 Components 表中定义的组件。"
                 "同一个组件类型可以被多次实例化（不同 PrototypeName）。")
    _add_section(doc, "基础数据类型", "PrimitiveDataTypes",
                 "定义 APPLICATION-PRIMITIVE-DATA-TYPE 及其 IMPLEMENTATION-DATA-TYPE 映射。"
                 "BaseType 选择平台支持的 boolean/uint8/uint16/uint32/uint64/sint8/sint16/sint32/float32。"
                 "CompuMethodRef 指向 CompuMethods 表，DataConstrRef 指向 DataConstrs 表，UnitRef 指向 Units 表。")
    _add_section(doc, "结构体类型", "RecordTypes",
                 "APPLICATION-RECORD-DATA-TYPE 定义。结构体对应的 IMPLEMENTATION-DATA-TYPE 由工具自动生成。")
    _add_section(doc, "结构体字段", "RecordElements",
                 "Record 类型的字段定义。Order 决定字段在结构体中的排列顺序（从 1 开始）。")
    _add_section(doc, "数据类型映射", "DataTypeMappings",
                 "显式定义 ADT 到 IDT 的映射关系。所有映射归入同一个 DATA-TYPE-MAPPING-SET。"
                 "每个 SWC 的 InternalBehavior 会自动引用该 MappingSet。")
    _add_section(doc, "计算方法", "CompuMethods",
                 "COMPU-METHOD 定义。Category 可选：TEXTTABLE（文本映射表）、LINEAR（线性缩放）、IDENTICAL（直接传递）。")
    _add_section(doc, "计算标度", "CompuScales",
                 "CompuMethod 的标度/映射条目。TEXTTABLE 填写 LowerLimit/UpperLimit/TextValue；"
                 "LINEAR 填写 Numerator/Denominator/Offset。")
    _add_section(doc, "数据约束", "DataConstrs",
                 "DATA-CONSTR 定义。LowerLimit/UpperLimit 指定数据范围（闭区间）。")
    _add_section(doc, "SR 接口", "SRInterfaces",
                 "SENDER-RECEIVER-INTERFACE 定义。IsService=false 表示非 Service 接口。"
                 "数据元素在 SRDataElements 表中定义。")
    _add_section(doc, "SR 数据元素", "SRDataElements",
                 "S/R 接口的数据元素。每个元素引用一个 ApplicationDataType。")
    _add_section(doc, "CS 接口", "CSInterfaces",
                 "CLIENT-SERVER-INTERFACE 定义。Operation 在 CSOperations 表中定义，"
                 "参数在 CSArguments 表中定义。")
    _add_section(doc, "CS 操作", "CSOperations",
                 "C/S 接口的 Operation 声明。")
    _add_section(doc, "CS 参数", "CSArguments",
                 "C/S Operation 的参数定义。Direction 可选 IN/OUT/INOUT。"
                 "每个参数引用一个 ApplicationDataType。")
    _add_section(doc, "端口", "Ports",
                 "SWC 端口定义。PortDirection: P=Provider/Server, R=Requester/Client。"
                 "InterfaceKind: SR 或 CS。InterfaceRef 为接口的完整路径。"
                 "ComSpecKind 根据方向和接口类型选择对应的 COM-SPEC 类型。"
                 "AliveTimeout/QueueLength/EnableUpdate/HandleTimeoutType/InitValue 根据 ComSpec 类型填写。")
    _add_section(doc, "Runnables", "Runnables",
                 "Runnable 实体定义。Symbol 是生成的 C 函数名。")
    _add_section(doc, "Runnable 事件", "RunnableEvents",
                 "Runnable 的触发事件。TriggerType: Init/Periodic/OperationInvoked/DataReceived。"
                 "Periodic 需填 PeriodMs；OperationInvoked 需填 PortName+OperationName；"
                 "DataReceived 需填 PortName+DataElementName。")
    _add_section(doc, "Runnable 访问点", "RunnableAccesses",
                 "Runnable 内的数据/服务访问点。AccessType: DataRead（读 SR R-Port 数据）、"
                 "DataWrite（写 SR P-Port 数据）、ServerCallPoint（调用 CS R-Port 服务）。"
                 "AccessName 用于生成 ACCESS-POINT 的 SHORT-NAME。")
    _add_section(doc, "组合连接器", "CompositionConnectors",
                 "Composition 内的 Assembly 连接器。ConnectorType 固定为 Assembly。"
                 "ProviderPrototype/RequesterPrototype 引用 ComponentPrototypes 表中的实例名。"
                 "ProviderPort/RequesterPort 引用 Ports 表中的端口名。")
    _add_section(doc, "物理单位", "Units",
                 "UNIT 定义。FactorSIToUnit 和 OffsetSIToUnit 描述 SI 单位转换关系。"
                 "无单位类型使用 Unit_NoUnit（因子 1，偏移 0）。")

    # ── Appendix: naming conventions ──
    doc.add_heading("附录：命名规范", level=1)
    doc.add_paragraph("端口命名：Pp_<组件>_<接口类型>（P-Port），Rp_<目标组件>_<接口类型>（R-Port）")
    doc.add_paragraph("接口命名：If_<信号名>_SR（S/R 接口），If_<服务名>_CS（C/S 接口）")
    doc.add_paragraph("Operation 命名：rr<动作>（Requester 调用），get<状态>（查询状态）")
    doc.add_paragraph("Runnable 命名：<SWC名>_Init / _Step / _<功能>")
    doc.add_paragraph("ServerCallPoint 命名：SC_<R-Port名>_<Operation名>")
    doc.add_paragraph("DataAccess 命名：DR_<Port名>_<DataElement>（读）/ DW_<Port名>_<DataElement>（写）")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"DOCX template created: {output}")


def _add_section(doc, title_cn: str, sheet_name: str, description: str) -> None:
    """Add a document section with heading, description, and data table."""
    doc.add_heading(f"{title_cn}（{sheet_name}）", level=2)
    doc.add_paragraph(description)

    headers = SHEETS.get(sheet_name, [])
    rows = EXAMPLE_ROWS.get(sheet_name, [])

    if not headers:
        doc.add_paragraph("（无示例数据）")
        return

    # Create table: 1 header + N data rows
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_shading(hdr_cells[i], "4472C4")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
                run.font.bold = True

    # Data rows
    for r, row_data in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, val in enumerate(row_data):
            if c < len(headers):
                row_cells[c].text = str(val) if val else ""
                for p in row_cells[c].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    doc.add_paragraph()  # spacer


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading_elm = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    tcPr.append(shading_elm)


# ── DOCX → Excel conversion ──────────────────────────────────────

def docx_to_excel(docx_path: Path, excel_path: Path) -> None:
    """Parse a structured DOCX and produce an Excel workbook.

    The DOCX is expected to follow the format produced by create_docx_template():
    - Each sheet is a section with a heading containing the sheet name in parentheses.
    - A table follows immediately after the heading.
    - The table's first row is the header (must match the expected column names).
    - Subsequent rows are data.
    """
    doc = Document(str(docx_path))

    # Build sheet name → headers mapping
    sheet_headers = {name: [h.lower().strip() for h in headers] for name, headers in SHEETS.items()}

    # Parse document using python-docx iterator
    workbook = Workbook()
    workbook.remove(workbook.active)

    current_sheet: str | None = None
    header_fill = PatternFill("solid", fgColor="E2F0D9")
    processed_sheets: set[str] = set()

    # Iterate through block-level elements in document order
    body = doc.element.body
    nsmap_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # Only match Heading 2 paragraphs (not TOC entries)
            if not _is_heading2(child, nsmap_w):
                continue
            text = _extract_text(child, nsmap_w)
            if text:
                sheet = _match_sheet_name(text, sheet_headers)
                if sheet and sheet not in processed_sheets:
                    current_sheet = sheet

        elif tag == "tbl":
            if current_sheet:
                _write_table_to_sheet(workbook, current_sheet, child, header_fill)
                processed_sheets.add(current_sheet)
                current_sheet = None

    # Add validation
    _add_excel_validations(workbook)

    excel_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(str(excel_path))
    print(f"Excel created: {excel_path} (sheets: {len(processed_sheets)})")


def _is_heading2(p_element, ns_w: str) -> bool:
    """Check if a paragraph has Heading 2 style."""
    for pPr in p_element.iter(f"{{{ns_w}}}pPr"):
        for pStyle in pPr.iter(f"{{{ns_w}}}pStyle"):
            val = pStyle.get(f"{{{ns_w}}}val", "")
            if val in ("Heading2", "2"):
                return True
    return False


def _extract_text(element, ns_w: str) -> str:
    """Extract all text from a paragraph element."""
    parts = []
    for t in element.iter(f"{{{ns_w}}}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts).strip()


def _match_sheet_name(text: str, sheet_headers: dict[str, list[str]]) -> str | None:
    """Check if the heading text references a known sheet name."""
    for name in sheet_headers:
        if name.lower() in text.lower():
            return name
    return None


def _write_table_to_sheet(workbook, sheet_name: str, tbl, header_fill) -> None:
    """Extract table data and write to an Excel sheet."""
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rows_data = []
    for tr in tbl.iter(f"{{{ns_w}}}tr"):
        row = []
        for tc in tr.iter(f"{{{ns_w}}}tc"):
            texts = []
            for p in tc.iter(f"{{{ns_w}}}p"):
                texts.append(_extract_text(p, ns_w))
            row.append("".join(texts).strip())
        # Skip completely empty rows
        if any(cell for cell in row):
            rows_data.append(row)

    if not rows_data:
        return

    sheet = workbook.create_sheet(sheet_name)

    # Write header row
    for c, val in enumerate(rows_data[0]):
        cell = sheet.cell(row=1, column=c + 1, value=val)
        cell.font = XlFont(bold=True)
        cell.fill = header_fill

    # Write data rows
    for r, row in enumerate(rows_data[1:], start=2):
        for c, val in enumerate(row):
            sheet.cell(row=r, column=c + 1, value=val)

    sheet.freeze_panes = "A2"

    # Auto-fit column width
    for col_idx, col_cells in enumerate(sheet.columns, 1):
        max_width = 14
        for cell in col_cells:
            if cell.value:
                max_width = max(max_width, min(len(str(cell.value)) + 2, 64))
        sheet.column_dimensions[sheet.cell(row=1, column=col_idx).column_letter].width = max_width


def _add_excel_validations(workbook: Workbook) -> None:
    """Add dropdown validations to the generated Excel."""
    rules = {
        "Components": {"B": '"Application,Composition"'},
        "PrimitiveDataTypes": {"E": '"boolean,uint8,uint16,uint32,uint64,sint8,sint16,sint32,float32"', "H": '"READ-ONLY,READ-WRITE,NOT-ACCESSIBLE"'},
        "CompuMethods": {"C": '"TEXTTABLE,LINEAR,IDENTICAL"'},
        "SRInterfaces": {"C": '"true,false"'},
        "CSInterfaces": {"C": '"true,false"'},
        "CSArguments": {"D": '"IN,OUT,INOUT"'},
        "Ports": {"C": '"P,R"', "D": '"SR,CS"', "H": '"CLIENT-COM-SPEC,SERVER-COM-SPEC,NONQUEUED-SENDER-COM-SPEC,NONQUEUED-RECEIVER-COM-SPEC,QUEUED-SENDER-COM-SPEC,QUEUED-RECEIVER-COM-SPEC"'},
        "RunnableEvents": {"C": '"Init,Periodic,OperationInvoked,DataReceived"'},
        "RunnableAccesses": {"C": '"DataRead,DataWrite,ServerCallPoint"'},
        "CompositionConnectors": {"F": '"Assembly"'},
    }
    for sheet_name, column_rules in rules.items():
        if sheet_name not in workbook.sheetnames:
            continue
        sheet = workbook[sheet_name]
        for column, formula in column_rules.items():
            validation = DataValidation(type="list", formula1=formula, allow_blank=True)
            sheet.add_data_validation(validation)
            validation.add(f"{column}2:{column}1000")


# ── CLI ───────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="DOCX ↔ Excel converter for ARXML codegen")
    sub = parser.add_subparsers(dest="command", required=True)

    c1 = sub.add_parser("create-docx", help="Generate standard DOCX template")
    c1.add_argument("--output", type=Path, default=Path("docs/horn_standard.docx"), help="Output DOCX path")

    c2 = sub.add_parser("convert", help="Convert DOCX to Excel")
    c2.add_argument("--input", type=Path, required=True, help="Input DOCX path")
    c2.add_argument("--output", type=Path, required=True, help="Output Excel path")

    args = parser.parse_args()

    if args.command == "create-docx":
        create_docx_template(args.output)
    elif args.command == "convert":
        docx_to_excel(args.input, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
